"""
Document indexer — parse PDF/DOCX/TXT files into chunks and index into Pinecone.

Multi-tenant security model: same as rag_service.py
  Namespace  = tenant_id       → physical company wall in Pinecone
  Metadata   = doc_id + user_id → row-level isolation

Chunk strategy:
  - Paragraph-level chunking with 800-char target, 100-char overlap
  - Metadata carries doc_id, page, chunk_index, source_type="document"
  - Titles/headings from DOCX structure are prepended to chunks for richer embedding
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

log = logging.getLogger(__name__)

CHUNK_SIZE    = 800   # target chars per chunk
CHUNK_OVERLAP = 100   # chars of overlap between adjacent chunks
UPSERT_BATCH  = 100
EMBED_BATCH   = 96
EMBED_MODEL   = "voyage-3"


# ── Text extraction ────────────────────────────────────────────────────────────

def _extract_txt(path: str) -> list[dict]:
    """Returns list of {page: int, text: str}."""
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    return [{"page": 1, "text": text}]


def _extract_pdf(path: str) -> list[dict]:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append({"page": i, "text": text})
    return pages


def _extract_docx(path: str) -> list[dict]:
    from docx import Document
    doc = Document(path)
    pages: list[dict] = []
    current_lines: list[str] = []
    virtual_page = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Treat each paragraph as having a style tag for context
        style = para.style.name if para.style else ""
        if "Heading" in style:
            # Flush current page before new section
            if current_lines:
                pages.append({"page": virtual_page, "text": "\n".join(current_lines)})
                current_lines = []
                virtual_page += 1
            current_lines.append(f"## {text}")
        else:
            current_lines.append(text)
            # Soft page boundary every ~50 paragraphs
            if len(current_lines) >= 50:
                pages.append({"page": virtual_page, "text": "\n".join(current_lines)})
                current_lines = []
                virtual_page += 1

    if current_lines:
        pages.append({"page": virtual_page, "text": "\n".join(current_lines)})

    return pages


def extract_pages(file_path: str, file_type: str) -> list[dict]:
    """Dispatch to format-specific extractor. Returns [{page, text}]."""
    t = file_type.lower().lstrip(".")
    try:
        if t == "pdf":
            return _extract_pdf(file_path)
        if t in ("docx", "doc"):
            return _extract_docx(file_path)
        return _extract_txt(file_path)
    except Exception as exc:
        log.error("Text extraction failed (%s, %s): %s", file_path, file_type, exc)
        raise


# ── Chunking ───────────────────────────────────────────────────────────────────

def _split_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping character-level chunks, preferring paragraph breaks."""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= size:
        return [text] if text else []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + size
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        # Try to break at a paragraph or sentence boundary within the last 200 chars
        window = text[start:end]
        break_pos = max(
            window.rfind("\n\n"),
            window.rfind(". "),
            window.rfind("? "),
            window.rfind("! "),
        )
        if break_pos > size // 2:
            end = start + break_pos + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
    return chunks


def build_chunks(
    doc_id:    str,
    tenant_id: str,
    user_id:   str,
    doc_name:  str,
    pages:     list[dict],
) -> list[dict]:
    """
    Convert extracted pages into indexable chunks.
    Returns [{"id", "text", "metadata"}].
    """
    import time as _time
    ts = _time.strftime("%Y-%m-%dT%H:%M:%SZ", _time.gmtime())
    chunks: list[dict] = []
    chunk_index = 0

    for page_info in pages:
        page_num = page_info["page"]
        page_text = page_info["text"].strip()
        if not page_text:
            continue

        for segment in _split_text(page_text):
            vec_id = f"{tenant_id}_{doc_id}_chunk_{chunk_index}"
            chunks.append({
                "id": vec_id,
                "text": segment,
                "metadata": {
                    "text":         segment,
                    "doc_id":       doc_id,
                    "tenant_id":    tenant_id,
                    "user_id":      user_id,
                    "document_title": doc_name,
                    "source_type":  "document",
                    "chunk_type":   "document_chunk",
                    "page":         page_num,
                    "chunk_index":  chunk_index,
                    "timestamp":    ts,
                },
            })
            chunk_index += 1

    return chunks


# ── Main indexer class ─────────────────────────────────────────────────────────

class DocumentIndexer:
    """
    Lazy-initialised. Shares Voyage + Pinecone instances with RAGService
    but operates independently for document-specific indexing.
    """

    def __init__(self):
        self._voyage = None
        self._index  = None
        self._ready: bool | None = None

    def _init(self) -> bool:
        if self._ready is not None:
            return self._ready
        from backend.config import settings

        missing = []
        if not settings.voyageai_api_key:
            missing.append("VOYAGEAI_API_KEY")
        if not settings.pinecone_api_key:
            missing.append("PINECONE_API_KEY")
        if not settings.pinecone_index:
            missing.append("PINECONE_INDEX")
        if missing:
            log.warning("DocumentIndexer disabled — missing: %s", ", ".join(missing))
            self._ready = False
            return False

        try:
            import voyageai
            self._voyage = voyageai.Client(api_key=settings.voyageai_api_key)
        except ImportError:
            log.warning("DocumentIndexer disabled — install voyageai")
            self._ready = False
            return False

        try:
            from pinecone import Pinecone
            self._index = Pinecone(api_key=settings.pinecone_api_key).Index(settings.pinecone_index)
        except Exception as exc:
            log.warning("DocumentIndexer disabled — Pinecone: %s", exc)
            self._ready = False
            return False

        self._ready = True
        return True

    def _embed(self, texts: list[str]) -> list[list[float]]:
        import time as _time
        all_emb: list[list[float]] = []
        for i in range(0, len(texts), EMBED_BATCH):
            batch = texts[i: i + EMBED_BATCH]
            for attempt in range(3):
                try:
                    result = self._voyage.embed(batch, model=EMBED_MODEL, input_type="document")
                    all_emb.extend(result.embeddings)
                    break
                except Exception as exc:
                    err = str(exc)
                    is_rate = "rate" in err.lower() or "429" in err or "too many" in err.lower()
                    if is_rate and attempt < 2:
                        _time.sleep(25 * (attempt + 1))
                        continue
                    raise
        return all_emb

    def index_document(
        self,
        doc_id:    str,
        tenant_id: str,
        user_id:   str,
        file_path: str,
        file_type: str,
        doc_name:  str,
    ) -> int:
        """
        Parse file → chunk → embed → upsert to Pinecone namespace=tenant_id.
        Returns number of vectors indexed (0 on failure).
        Raises on hard errors so the caller can mark status=FAILED.
        """
        if not self._init():
            log.warning("DocumentIndexer not ready — vectors not indexed for doc %s", doc_id)
            return 0

        pages  = extract_pages(file_path, file_type)
        chunks = build_chunks(doc_id, tenant_id, user_id, doc_name, pages)
        if not chunks:
            log.info("DocumentIndexer: no chunks extracted from %s", doc_id)
            return 0

        texts      = [c["text"] for c in chunks]
        embeddings = self._embed(texts)

        vectors = [
            {"id": c["id"], "values": emb, "metadata": c["metadata"]}
            for c, emb in zip(chunks, embeddings)
        ]

        for i in range(0, len(vectors), UPSERT_BATCH):
            self._index.upsert(
                vectors=vectors[i: i + UPSERT_BATCH],
                namespace=tenant_id,
            )

        log.info("DocumentIndexer: %d vectors indexed (doc=%s tenant=%s)", len(vectors), doc_id, tenant_id)
        return len(vectors)

    def delete_document(self, doc_id: str, tenant_id: str) -> None:
        """Remove all vectors for a document from the tenant namespace."""
        if not self._init():
            return
        try:
            prefix  = f"{tenant_id}_{doc_id}_chunk_"
            id_list: list[str] = []
            for page in self._index.list(prefix=prefix, namespace=tenant_id):
                for item in (page.vectors or []):
                    id_list.append(item.id)
            if id_list:
                self._index.delete(ids=id_list, namespace=tenant_id)
                log.info("DocumentIndexer: deleted %d vectors (doc=%s)", len(id_list), doc_id)
        except Exception as exc:
            log.warning("DocumentIndexer delete failed (doc=%s): %s", doc_id, exc)

    def page_count(self, file_path: str, file_type: str) -> int:
        """Return page count without full extraction (fast path for metadata)."""
        try:
            t = file_type.lower().lstrip(".")
            if t == "pdf":
                from pypdf import PdfReader
                return len(PdfReader(file_path).pages)
            if t in ("docx", "doc"):
                from docx import Document
                doc = Document(file_path)
                # Approximate: count paragraphs / 40 as proxy for pages
                paras = sum(1 for p in doc.paragraphs if p.text.strip())
                return max(1, paras // 40)
            # TXT: lines / 50
            lines = Path(file_path).read_text(encoding="utf-8", errors="replace").count("\n")
            return max(1, lines // 50)
        except Exception:
            return 1


# Module singleton
doc_indexer = DocumentIndexer()

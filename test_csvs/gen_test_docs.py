"""Generate test documents (TXT, DOCX, PDF) for Playwright QA."""
import pathlib

OUT = pathlib.Path(__file__).parent / "test_docs"
OUT.mkdir(exist_ok=True)

# ── TXT (small, 1 page equivalent) ──────────────────────────────────────────
txt = OUT / "small_report.txt"
txt.write_text("""\
Quarterly Demand Report — Q1 2026
==================================

Executive Summary
-----------------
Total demand for Q1 2026 was 8.3% above forecast.
Key drivers: Semana Santa (Colombia) and Medellin expansion.

SKU Performance
---------------
SKU-001 (Premium Widget): 12,450 units sold. Forecast was 11,200. WAPE=8.5%.
SKU-002 (Standard Widget): 8,920 units sold. Forecast was 9,100. WAPE=3.2%.
SKU-003 (Economy Widget): 22,100 units sold. Forecast was 19,800. WAPE=11.6%.

Inventory Recommendations
-------------------------
- Increase reorder point for SKU-001 by 15% due to volatility (CV=0.34)
- Safety stock for SKU-003 is critically low — replenish within 7 days
- SKU-002 is healthy, no action required

Model Comparison
----------------
LightGBM won on 7 of 10 SKUs. Average WAPE: 9.2%.
Prophet was best for SKU-007 and SKU-009 (strong weekly seasonality).

Action Items for Q2
-------------------
1. Update promotion calendar in the config
2. Retrain models after April holiday data is collected
3. Review SKU-005 for intermittent demand treatment (Croston recommended)
""", encoding="utf-8")
print(f"Created: {txt}")

# ── DOCX (medium, with headings) ─────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Forecasting Platform — User Guide", 0)

    doc.add_heading("1. Getting Started", 1)
    doc.add_paragraph(
        "This platform uses machine learning models to generate demand forecasts "
        "for your product portfolio. Supported models include LightGBM, XGBoost, "
        "ARIMA, Prophet, ETS, and Croston. Each model is trained per SKU with "
        "walk-forward cross-validation to prevent data leakage."
    )

    doc.add_heading("2. Uploading Data", 1)
    doc.add_paragraph(
        "Navigate to the Data page and upload a CSV, Excel, or Parquet file. "
        "The file must contain at minimum: a date column, a target (sales/demand) "
        "column, and optionally a SKU/group identifier column."
    )

    doc.add_heading("3. Configuring a Forecast Session", 1)
    doc.add_paragraph(
        "Use the Forecast Studio wizard to configure column mappings, feature "
        "engineering (lags, rolling windows, calendar features), model selection, "
        "and validation strategy. The wizard has 8 steps."
    )

    doc.add_heading("4. Understanding Results", 1)
    doc.add_paragraph(
        "After training completes, the Results step shows MAE, RMSE, and WAPE "
        "for each model-SKU combination. Lower WAPE is better. A WAPE below 10% "
        "is considered excellent for consumer goods demand forecasting."
    )

    doc.add_heading("5. AI Analyst", 1)
    doc.add_paragraph(
        "The Agent AI page allows you to ask natural language questions about your "
        "forecasts. Example queries:\n"
        "- Which SKU has the worst forecast accuracy?\n"
        "- What is the reorder point for SKU-001?\n"
        "- Compare LightGBM vs Prophet across all SKUs"
    )

    doc.add_heading("6. Inventory Recommendations", 1)
    doc.add_paragraph(
        "The platform automatically calculates safety stock and reorder points "
        "based on forecast error and lead time. Stockout risk is shown as a "
        "percentage. SKUs above 30% risk are flagged for immediate action."
    )

    for _ in range(5):
        doc.add_heading("Appendix: Glossary", 2)
        doc.add_paragraph(
            "WAPE (Weighted Absolute Percentage Error): Sum of absolute errors "
            "divided by sum of actuals. Preferred over MAPE for zero-heavy series.\n"
            "MAE (Mean Absolute Error): Average of absolute forecast errors.\n"
            "RMSE (Root Mean Squared Error): Square root of mean squared errors.\n"
            "Walk-forward validation: Time-series cross-validation that expands "
            "the training window without allowing future data to leak into training."
        )

    docx_path = OUT / "user_guide.docx"
    doc.save(str(docx_path))
    print(f"Created: {docx_path}")
except ImportError:
    print("python-docx not installed — skipping DOCX")

# ── PDF (using reportlab, already installed) ──────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    pdf_path = OUT / "technical_spec.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("ForecastingCore — Technical Specification", styles["Title"]))
    story.append(Spacer(1, 0.5*cm))

    sections = [
        ("Pipeline Overview",
         "ForecastingCore is a Python ML library that orchestrates multi-model "
         "time-series forecasting. The pipeline consists of: data loading, "
         "feature engineering, model training (per SKU), evaluation, and forecast "
         "generation. All models are sklearn-compatible and support walk-forward "
         "cross-validation."),
        ("Feature Engineering",
         "Calendar features include day-of-week, month, week-of-year, and Colombian "
         "holiday distances (Easter, Christmas, Semana Santa). Lag features cover "
         "configurable look-back windows. Rolling statistics include mean, std, min, "
         "max and trend via polynomial fit."),
        ("ML Models (LightGBM / XGBoost)",
         "Gradient boosting models trained per SKU. Hyperparameter tuning via Optuna "
         "(optional). Walk-forward validation with expanding window. Residuals stored "
         "for prediction interval construction at inference time."),
        ("Statistical Models (ARIMA, Prophet, ETS, Croston)",
         "Statistical models bypass feature engineering and operate on raw time series. "
         "ARIMA supports seasonal differencing. Prophet handles multiple seasonalities "
         "and holiday effects. ETS covers trend/seasonal components. Croston targets "
         "intermittent demand patterns."),
        ("Evaluation Metrics",
         "MAE: mean absolute error. RMSE: root mean squared error. WAPE: weighted "
         "absolute percentage error (robust to zeroes). Bias: directional error "
         "indicating systematic over/under-forecasting. All metrics computed "
         "out-of-sample on the held-out test fold."),
        ("Inventory Advisor",
         "Safety stock = Z * sigma * sqrt(lead_time). Reorder point = mu * lead_time + "
         "safety_stock. Stockout risk estimated from forecast error distribution. "
         "Holding cost proportional to average inventory level."),
    ]

    for title, body in sections:
        story.append(Paragraph(title, styles["Heading2"]))
        story.append(Paragraph(body, styles["BodyText"]))
        story.append(Spacer(1, 0.4*cm))

    doc.build(story)
    print(f"Created: {pdf_path}")
except ImportError:
    print("reportlab not installed — skipping PDF")

print("\nAll test documents ready in:", OUT)
